import os
from docx import Document
import markdown
import webbrowser
from datetime import datetime

# from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.style import WD_BUILTIN_STYLE
from docx.shared import Pt

from data_sets import head_sets, head_sets1, head_sets2


def add_lead_marks(count, mark='#'):
    prefix = '#'
    for level_count in range(0, count):
        prefix = prefix+mark
    return prefix


def prep_markdown(head_list):
    markdown_strings = []
    for head in head_list:
        mrkdwn = {}
        level = add_lead_marks(head['heading_level'])
        hstring = level+' '+head['title']
        mrkdwn[head['id']] = {'hstring':hstring}
        markdown_strings.append(mrkdwn)
        print(hstring)
        # print(hstring.replace('#', ' '))
    return markdown_strings


def prep_html_from_markdown(head_list):
    html = ''
    for r in head_list:
        for k, v in r.items():
            n = r[k]['hstring'] + '\n'
            # print(n)
        html += n
    html = markdown.markdown(html)
    # print(html)
    return html


def prep_html_raw(head_list):
    modlist = ''
    for head in head_list:
        level = head['heading_level']+1
        hstring = f" <h{level} id={head['id']} > {head['title']} </h{level}>\n"
        modlist += hstring
        # print(hstring)
    return modlist


def save_to_html(fname, cont):
    try:
        with open(fname, "w", encoding="utf-8", errors="xmlcharrefreplace") as output_file:
            output_file.write(cont)
    except FileExistsError:
        fname = fname+'_'+datetime.now().strftime('%Y-%m-%d_%H-%m')+'.html'
        with open(fname,"w", encoding="utf-8",
                  errors="xmlcharrefreplace") as output_file:
            output_file.write(cont)
    return 'file://' + os.path.realpath(fname)

def prep_docx_from_template(head_list):
    """
    Seems that autonumeration doesn't work wit document created directly from python-docx
    :param head_list:
    :return:
    """
    document = Document('template.docx')
    sections = document.sections
    section = sections[0]
    section.style = 'Heading 1'
    # styles = document.styles
    headings_types = {}
    for x in range(0, 10):
        headings_types[x] = 0
    print(headings_types)
    for head in head_list:
        print(head)
        # if head['heading_level'] > 9: #max heading for docx is 9
        #     head['heading_level'] = 9
        #     print(f"Decreased indentation level to supported value")
        headings_types[head['heading_level']] += 1
        print(head['heading_level'], headings_types[head['heading_level']])

        document.add_heading(f"{head['title']}", level=head['heading_level']+1) #False heading level caused by lack of knowledge about styles apply 'Title' style handling
        # print(f"headlevel = {head['heading_level']}")

        # headings_types[head['heading_level']] = headings_types[head['heading_level']]+1
    print(headings_types)
    try:
        document.save('templated.docx')
    except PermissionError:
        document.save('templated_'+datetime.now().strftime('%Y-%m-%d_%H-%m')+'.docx')


def prep_docx(head_list):
    document = Document()
    paragraph = document.add_paragraph(style='Normal')
    paragraph.style = document.styles['Heading 1']

    headings_types = {}

    for x in range(0,10): headings_types[x] = 0
    print(headings_types)
    for head in head_list:
        print(head)
        if head['heading_level'] > 9: #max heading for docx is 9
            head['heading_level'] = 9
            print(f"Decreased indentation level to supported value")

        headings_types[head['heading_level']] += 1
        print(head['heading_level'], headings_types[head['heading_level']])
        mtab = head['heading_level']*'\t'

        numeration = f"{'1.'*(head['heading_level']-1)}{headings_types[head['heading_level']]} "
        document.add_heading(f"{mtab}{numeration}{head['title']}", head['heading_level']+1)

    print(headings_types)
    try:
        document.save('testowy.docx')
    except PermissionError:
        document.save('testowy_'+datetime.now().strftime('%Y-%m-%d_%H-%m')+'.docx')
def data_set_validation(dset):
    data_correct = True
    # print(f"DATA SET TYPE: {type(dset)}")
    data_correct = data_correct and isinstance(dset, list)
    # print(f"DATA CORRECT: {data_correct}")
    for item in dset:
        # print(f"DATA SUBSET TYPE: {type(item)}")
        data_correct = data_correct and isinstance(item, dict)

        # print(f"DATA CORRECT: {data_correct}")
        # print(dset[0].keys())
        data_correct = data_correct and 'id' in item.keys()
        # print(f"DATA CORRECT: {data_correct}")
        data_correct = data_correct and 'title' in item.keys()
        # print(f"DATA CORRECT: {data_correct}")
        data_correct = data_correct and 'heading_level' in item.keys()
        # print(f"DATA CORRECT: {data_correct}")
        data_correct = data_correct and isinstance(item['id'], int)
        # print(f"DATA CORRECT: {data_correct}")
        data_correct = data_correct and isinstance(item['title'], str)
        # print(f"DATA CORRECT: {data_correct}")
        data_correct = data_correct and isinstance(item['heading_level'], int)
        # print(f"DATA CORRECT: {data_correct}")
    print(f"DATA CORRECT: {data_correct}")
    return data_correct

def main():
    hset = head_sets1
    if data_set_validation(hset):

        prep_docx_from_template(hset)
        res = prep_markdown(hset)
        print(f"MARKDOWN STRINGS:\n{res}")
        raw_html = prep_html_raw(hset)
        print(f"RAW HTML:\n{raw_html}")
        html = prep_html_from_markdown(res)
        print(f"HTML:\n{html}")
        prep_docx(hset)

        html_file= save_to_html('raw_html.html', raw_html)
        webbrowser.open(html_file)

        # html = prep_html(res)
        # save_to_html('test.html', html)
    else:
        print("Something wrong with input data")


if __name__ == '__main__':
    main()